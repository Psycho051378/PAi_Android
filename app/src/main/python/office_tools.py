"""
Office Tools — Word / Excel / PowerPoint via Python.

Usage (JSON over stdin/sys.argv):
  python office_tools.py '{"action":"word_create","data":{...}}'

Returns JSON: {"status":"ok","path":"...","preview":"...","error":null}

Supported actions:
  word_create   — create .docx from JSON template
  word_read     — read .docx, return text/tables
  excel_create  — create .xlsx from JSON data
  excel_read    — read .xlsx, return rows
  pptx_create   — create .pptx from JSON (basic, no images)
"""

import json, sys, os, datetime, io, zipfile, xml.etree.ElementTree as ET

# ──────────────── Word (python-docx) ────────────────

def word_create(data):
    """Create a .docx file from JSON template.
    data format:
    {
        "path": "/path/to/output.docx",
        "title": "Document Title",
        "heading1": "Section 1",
        "paragraphs": ["Text1", "Text2"],
        "tables": [
            {"rows": 3, "cols": 4, "header": ["H1","H2","H3","H4"],
             "cells": [["A1","B1","C1","D1"], ["A2","B2","C2","D2"]]}
        ],
        "bold": ["text to be bold"],
        "italic": ["text to be italic"]
    }
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return err("python-docx not installed")

    doc = Document()
    path = data.get("path", "")

    # Default styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = data.get("title", "")
    if title:
        p = doc.add_heading(title, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading 1
    h1 = data.get("heading1", "")
    if h1:
        doc.add_heading(h1, 1)

    # Heading 2
    h2 = data.get("heading2", "")
    if h2:
        doc.add_heading(h2, 2)

    # Paragraphs
    for para_text in data.get("paragraphs", []):
        p = doc.add_paragraph(para_text)
        # Handle bold markers
        bold_texts = data.get("bold", [])
        for bt in bold_texts:
            if bt in para_text:
                p.clear()
                parts = para_text.split(bt)
                run = p.add_run(parts[0])
                run.bold = False
                run = p.add_run(bt)
                run.bold = True
                if len(parts) > 1:
                    run = p.add_run(parts[1])
                    run.bold = False

    # Tables
    for table_data in data.get("tables", []):
        rows = table_data.get("rows", 2)
        cols = table_data.get("cols", 2)
        header = table_data.get("header", [])
        cells = table_data.get("cells", [])

        if header:
            rows = max(rows, len(cells) + 1 if cells else rows)
            cols = max(cols, len(header))
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for j, h in enumerate(header):
                if j < cols:
                    cell = table.rows[0].cells[j]
                    cell.text = str(h) if h else ""
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            # Data rows
            for i, row_data in enumerate(cells):
                if i < rows - 1:
                    for j, val in enumerate(row_data):
                        if j < cols:
                            table.rows[i + 1].cells[j].text = str(val) if val else ""
        else:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            for i, row_data in enumerate(cells):
                if i < rows:
                    for j, val in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = str(val) if val else ""

        doc.add_paragraph()  # spacing after table

    # Bullet list
    for item in data.get("bullets", []):
        p = doc.add_paragraph(item, style='List Bullet')

    # Numbered list
    for item in data.get("numbered", []):
        p = doc.add_paragraph(item, style='List Number')

    # Save
    if not path:
        path = os.path.expanduser(f"~/document_{int(time.time())}.docx")

    doc.save(path)
    # Preview
    preview = data.get("title", "") or h1 or "Document"
    return ok(path, preview, f"Word document: {preview}")


def word_read(data):
    """Read .docx and return text + tables."""
    try:
        from docx import Document
    except ImportError:
        return err("python-docx not installed")

    path = data.get("path", "")
    if not path or not os.path.exists(path):
        return err("File not found: " + path)

    doc = Document(path)
    result_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            result_parts.append(para.text)

    tables_data = []
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            rows_data.append(row_data)
        tables_data.append(rows_data)

    full_text = "\n".join(result_parts)
    return ok(path, full_text[:500], full_text, {"tables": tables_data})


# ──────────────── Excel (openpyxl) ────────────────

def excel_create(data):
    """Create .xlsx file from JSON data.
    data format:
    {
        "path": "/path/to/output.xlsx",
        "sheets": [
            {
                "name": "Sheet1",
                "rows": [
                    ["Header1", "Header2", "Header3"],
                    ["val1", "val2", "val3"]
                ],
                "start_cell": "A1"
            }
        ]
    }
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        return err("openpyxl not installed")

    wb = Workbook()
    path = data.get("path", "")
    sheets_data = data.get("sheets", [])
    raw_query = data.get("raw_query", "")

    # Generate content from raw query if no structured data provided
    if not sheets_data and raw_query:
        import re
        # Try to extract column names from query (after ':' or after ',')
        col_match = re.search(r':\s*(.+?)$', raw_query)
        if col_match:
            cols_text = col_match.group(1)
            columns = [c.strip() for c in re.split(r'[,;]', cols_text) if c.strip()]
        else:
            # Fallback: extract key nouns
            columns = []
            for word in ['имя', 'возраст', 'город', 'name', 'age', 'city', 'цена', 'количество', 'сумма']:
                if word in raw_query.lower():
                    columns.append(word.capitalize())
            if not columns:
                columns = ['Колонка 1', 'Колонка 2', 'Колонка 3']

        # Generate sample rows
        sample_values = {
            'имя': ['Артём', 'Мария', 'Сергей', 'Анна', 'Дмитрий'],
            'name': ['John', 'Alice', 'Bob', 'Eve', 'Max'],
            'возраст': ['25', '34', '42', '29', '37'],
            'age': ['28', '35', '41', '22', '33'],
            'город': ['Москва', 'СПб', 'Казань', 'Екб', 'Нск'],
            'city': ['Moscow', 'SPb', 'Kazan', 'Ekb', 'Nsk'],
            'цена': ['100', '250', '500', '1000', '2000'],
            'количество': ['1', '2', '5', '10', '25'],
            'сумма': ['1000', '2500', '5000', '10000', '20000'],
        }

        rows = []
        rows.append(columns)
        for i in range(5):
            row = []
            for col in columns:
                col_lower = col.lower().strip()
                vals = sample_values.get(col_lower)
                if vals and i < len(vals):
                    row.append(vals[i])
                else:
                    row.append(f"{col}_{i+1}")
            rows.append(row)

        sheets_data = [{"name": "Sheet1", "rows": rows}]

    for idx, sheet_data in enumerate(sheets_data):
        name = sheet_data.get("name", f"Sheet{idx + 1}")
        if idx == 0:
            ws = wb.active
        else:
            ws = wb.create_sheet()
        ws.title = name

        rows = sheet_data.get("rows", [])
        start_cell = sheet_data.get("start_cell", "A1")

        # Parse start_cell (e.g., "B2" -> col=2, row=2)
        import re
        cell_match = re.match(r'([A-Z]+)(\d+)', start_cell)
        start_col = 1
        start_row = 1
        if cell_match:
            col_str = cell_match.group(1)
            start_row = int(cell_match.group(2))
            start_col = 0
            for c in col_str:
                start_col = start_col * 26 + (ord(c) - ord('A') + 1)

        for i, row_data in enumerate(rows):
            for j, val in enumerate(row_data):
                cell = ws.cell(row=start_row + i, column=start_col + j)
                cell.value = val

                # Header styling (first row)
                if i == 0:
                    cell.font = Font(bold=True, size=11)
                    cell.alignment = Alignment(horizontal='center')
                    cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

    # Auto-adjust column widths
    for ws in wb.worksheets:
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                try:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass
            adjusted_width = min(max_length + 3, 60)
            ws.column_dimensions[col_letter].width = adjusted_width

    if not path:
        path = os.path.expanduser(f"~/spreadsheet_{int(time.time())}.xlsx")

    wb.save(path)
    sheet_names = [s.get("name", f"Sheet{i+1}") for i, s in enumerate(sheets_data)]
    return ok(path, f"Sheets: {', '.join(sheet_names)}", f"Excel workbook with {len(sheets_data)} sheet(s)")


def excel_read(data):
    """Read .xlsx and return rows per sheet."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return err("openpyxl not installed")

    path = data.get("path", "")
    if not path or not os.path.exists(path):
        return err("File not found: " + path)

    wb = load_workbook(path, read_only=True, data_only=True)
    result = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data = []
        for row in ws.iter_rows(values_only=True):
            rows_data.append([str(v) if v is not None else "" for v in row])
        result[sheet_name] = {
            "rows": rows_data[:100],  # limit to 100 rows
            "total_rows": len(rows_data)
        }

    preview_lines = []
    for sname, sdata in result.items():
        if sdata["rows"]:
            preview_lines.append(f"Sheet '{sname}': {sdata['total_rows']} rows")
            preview_lines.append(" | ".join(sdata["rows"][0][:8]))

    wb.close()
    return ok(path, "\n".join(preview_lines[:20]), f"Read {len(result)} sheet(s)", {"sheets": result})


# ──────────────── PowerPoint (pure Python, no python-pptx) ────────────────

def pptx_create(data):
    """Create a .pptx using a pre-built template from python-pptx.
    The template is loaded from a file, slides duplicated and content modified.
    """
    try:
        import zipfile, os, re, shutil, io, xml.etree.ElementTree as ET
    except ImportError:
        return err("Required modules not available")

    path = data.get("path", "")
    template_path = data.get("template_path", "")
    slides_data = data.get("slides", [])
    raw_query = data.get("raw_query", "")

    topic_match = re.search(r'(?:про|о|об|на тему|по)\s+(.+?)(?:в формате|на\s+\d+|$)', raw_query, re.IGNORECASE)
    topic = topic_match.group(1).strip() if topic_match else data.get("title", "Новая презентация")
    slide_count = data.get("slideCount", 3)

    # Generate slide content
    slide_templates = [
        {"title": f"Введение в {topic}", "subtitle": "Обзор основных концепций",
         "bullets": [f"Что такое {topic}?", "История и предпосылки", "Основные понятия", "Актуальность"]},
        {"title": f"Ключевые аспекты {topic}", "subtitle": "Детальный разбор",
         "bullets": ["Первый ключевой аспект", "Второй ключевой аспект", "Третий ключевой аспект", "Взаимосвязь"]},
        {"title": f"Применение {topic}", "subtitle": "Практическое использование",
         "bullets": ["Современные применения", "Примеры из жизни", "Перспективы развития", "Выводы"]},
        {"title": f"Примеры {topic}", "subtitle": "Конкретные кейсы",
         "bullets": ["Пример 1", "Пример 2", "Анализ результатов", "Уроки"]},
        {"title": f"Будущее {topic}", "subtitle": "Тренды и прогнозы",
         "bullets": ["Текущие тренды", "Инновации", "Прогнозы", "Что дальше"]},
    ]
    slides_data = slide_templates[:min(slide_count, len(slide_templates))]
    if not slides_data:
        slides_data = [{"title": "New Presentation"}]

    if not template_path or not os.path.exists(template_path):
        return err("PPTX template not found at: " + str(template_path))

    # Copy template to output
    shutil.copy(template_path, path)

    # Read existing template files
    with zipfile.ZipFile(path, 'r') as zf:
        files = {name: bytearray(zf.read(name)) for name in zf.namelist()}

    # Get existing slide count from template
    existing_slides = sorted([n for n in files if n.startswith('ppt/slides/slide') and n.endswith('.xml')])

    # Decode and modify each slide
    for i, info in enumerate(slides_data):
        slide_name = f'ppt/slides/slide{i+1}.xml'
        if slide_name in files:
            slide_bytes = files[slide_name]
        elif existing_slides:
            slide_bytes = files[existing_slides[-1]]
        else:
            continue

        # Use string replacement to modify text content (preserves namespace prefixes)
        xml_str = slide_bytes.decode("utf-8")
        title = info.get("title", "Slide")
        subtitle = info.get("subtitle", "")
        bullets = info.get("bullets", [])

        # Find all a:t text elements with their content
        at_matches = list(re.finditer(r'<a:t[^>]*>.*?</a:t>', xml_str))

        if len(at_matches) >= 1:
            old_text = at_matches[0].group(0)
            xml_str = xml_str.replace(old_text, f'<a:t>{escape_xml(title)}</a:t>', 1)
        if len(at_matches) >= 2:
            old_text = at_matches[1].group(0)
            replacement = subtitle if subtitle else title
            xml_str = xml_str.replace(old_text, f'<a:t>{escape_xml(replacement)}</a:t>', 1)
        for j in range(2, min(len(at_matches), 2 + len(bullets))):
            if j - 2 < len(bullets):
                old_text = at_matches[j].group(0)
                xml_str = xml_str.replace(old_text, f'<a:t>{escape_xml(bullets[j - 2])}</a:t>', 1)

        files[slide_name] = xml_str.encode("utf-8")

    # Update presentation.xml slide count
    pres = files.get("ppt/presentation.xml", b"")
    if pres:
        pres_str = pres.decode("utf-8")
        # Update p:sldIdLst
        sld_pattern = r"<p:sldIdLst>.*?</p:sldIdLst>"
        new_slds = "<p:sldIdLst>\n"
        for i in range(len(slides_data)):
            new_slds += f'<p:sldId id="{i+1}" r:id="rId{i+3}"/>\n'
        new_slds += "</p:sldIdLst>"
        pres_str = re.sub(sld_pattern, new_slds, pres_str, flags=re.DOTALL)
        files["ppt/presentation.xml"] = pres_str.encode("utf-8")

    # Update pres_rels
    pres_rels = files.get("ppt/_rels/presentation.xml.rels", b"")
    if pres_rels:
        rels_str = pres_rels.decode("utf-8")
        # Remove old slide rels
        rels_str = re.sub(r'<Relationship Id="rId\d+" Type=".*?/slide" Target=".*?"/>\n?', "", rels_str)
        # Add new slide rels (starting from rId3)
        new_rels = ""
        for i in range(len(slides_data)):
            new_rels += f'<Relationship Id="rId{i+3}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i+1}.xml"/>\n'
        rels_str = rels_str.replace("</Relationships>", new_rels + "</Relationships>")
        rels_str = re.sub(r"\n\s*\n", "\n", rels_str)
        files["ppt/_rels/presentation.xml.rels"] = rels_str.encode("utf-8")

    # Remove extra slides from ZIP
    for i in range(len(slides_data), len(existing_slides)):
        for prefix in [f"ppt/slides/slide{i+1}.xml", f"ppt/slides/_rels/slide{i+1}.xml.rels"]:
            if prefix in files:
                del files[prefix]

    # Write back - use STORED (no compression) to avoid zlib issues on Android
    with zipfile.ZipFile(path, "w", zipfile.ZIP_STORED) as zf:
        for name, content in files.items():
            zf.writestr(name, bytes(content))

    return ok(path, f"Presentation with {len(slides_data)} slide(s)", f"Created: {len(slides_data)} slides, {os.path.getsize(path)} bytes")

def pptx_info(path):
    """Return basic info about a pptx file."""
    try:
        with zipfile.ZipFile(path, 'r') as zf:
            slides = [n for n in zf.namelist() if n.startswith('ppt/slides/slide') and n.endswith('.xml')]
            return f"{len(slides)} slides, {os.path.getsize(path)} bytes"
    except:
        return ""


# ──────────────── Format helpers ────────────────

def escape_xml(s):
    """Escape special XML characters."""
    if not s:
        return ""
    s = s.replace("&", "&amp;")
    s = s.replace("<", "&lt;")
    s = s.replace(">", "&gt;")
    s = s.replace('"', "&quot;")
    s = s.replace("'", "&apos;")
    return s


def ok(path, preview, description, extra=None):
    result = {"status": "ok", "path": path, "preview": preview, "description": description}
    if extra:
        result.update(extra)
    return json.dumps(result, ensure_ascii=False)


def err(msg):
    return json.dumps({"status": "error", "error": msg}, ensure_ascii=False)


# ──────────────── Main dispatcher ────────────────

def main():
    if len(sys.argv) < 2:
        print(err("Usage: office_tools.py <json_command>"))
        return

    raw = sys.argv[1]
    try:
        cmd = json.loads(raw)
    except json.JSONDecodeError as e:
        print(err(f"Invalid JSON: {e}"))
        return

    action = cmd.get("action", "")
    data = cmd.get("data", cmd)  # allow flat commands too

    handler = {
        "word_create": word_create,
        "word_read": word_read,
        "excel_create": excel_create,
        "excel_read": excel_read,
        "pptx_create": pptx_create,
    }

    fn = handler.get(action)
    if not fn:
        print(err(f"Unknown action: {action}. Available: {', '.join(handler.keys())}"))
        return

    try:
        result = fn(data)
        print(result)
    except Exception as e:
        import traceback
        print(err(f"{e}\n{traceback.format_exc()}"))


# main() is called by Android wrapper
def _entry_point():
    main()
